from dotenv import load_dotenv
load_dotenv()  # 👈 PHẢI GỌI TRƯỚC MỌI THỨ KHÁC
from openai import OpenAI


from fastapi import FastAPI, UploadFile, File, HTTPException
import os, shutil
import tiktoken
import json
import faiss
import numpy as np
from openai import OpenAI
import uuid
import traceback

import pdfplumber
from docx import Document

# =====================
# CONFIG
# =====================
VECTOR_DIM = 1536
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
VECTOR_DIR = os.path.join(BASE_DIR, "vector_db")
INDEX_PATH = os.path.join(VECTOR_DIR, "index.faiss")
META_PATH = os.path.join(VECTOR_DIR, "metadata.json")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(VECTOR_DIR, exist_ok=True)

client = OpenAI()

index = faiss.IndexFlatL2(VECTOR_DIM)
metadata = []
vector_ready = False


# =====================
# OPENAI CLIENT
# =====================


client = OpenAI()  # tự đọc OPENAI_API_KEY


# =====================
# TEXT EXTRACT
# =====================
def extract_text(path):
    if path.endswith(".pdf"):
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text

    if path.endswith(".docx"):
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    return ""


# =====================
# CHUNK TEXT
# =====================
def chunk_text(text, size=300):
    words = text.split()
    return [
        " ".join(words[i:i + size])
        for i in range(0, len(words), size)
    ]

# =====================
# EMBEDDING
# =====================
def embed(text: str):
    res = client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return res.data[0].embedding


# =====================
# BUILD VECTOR DB
# =====================
def build_vector_db():
    global metadata, vector_ready, index

    if vector_ready:
        return

    print("🔨 Building vector DB...")
    metadata = []
    index.reset()

    for filename in os.listdir(UPLOAD_DIR):
        path = os.path.join(UPLOAD_DIR, filename)
        text = extract_text(path)
        if not text.strip():
            continue

        chunks = chunk_text(text)
        for i, chunk in enumerate(chunks):
            vec = embed(chunk)
            index.add(np.array([vec]).astype("float32"))
            metadata.append({
                "content": chunk,
                "source": filename,
                "chunk_id": i
            })

    faiss.write_index(index, INDEX_PATH)
    json.dump(metadata, open(META_PATH, "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)

    vector_ready = True
    print(f"✅ Vector DB ready: {len(metadata)} chunks")


import os

from fastapi import FastAPI
from sqlalchemy import (
    create_engine,
    Column,
    Integer,
    String,
    Text,
    ForeignKey
)
from sqlalchemy.orm import (
    sessionmaker,
    declarative_base,
    relationship
)

# =====================
# APP
# =====================
app = FastAPI(title="Document Ingestion Service")

# =====================
# PATHS
# =====================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DB_PATH = os.path.join(BASE_DIR, "data.db")

os.makedirs(UPLOAD_DIR, exist_ok=True)

# =====================
# DATABASE
# =====================
engine = create_engine(
    f"sqlite:///{DB_PATH}",
    connect_args={"check_same_thread": False},
    echo=False  # đổi thành True nếu muốn debug SQL
)

SessionLocal = sessionmaker(
    autocommit=False,
    autoflush=False,
    bind=engine
)

Base = declarative_base()

# =====================
# MODELS
# =====================
class DocumentDB(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    original_filename = Column(String, nullable=False)
    stored_filename = Column(String, nullable=False)

    chunks = relationship(
        "ChunkDB",
        back_populates="document",
        cascade="all, delete-orphan"
    )


class ChunkDB(Base):
    __tablename__ = "chunks"

    id = Column(Integer, primary_key=True, index=True)
    doc_id = Column(Integer, ForeignKey("documents.id"), nullable=False)

    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)

    page = Column(Integer, nullable=True)
    section = Column(String, nullable=True)
    token_count = Column(Integer, nullable=False)

    document = relationship(
        "DocumentDB",
        back_populates="chunks"
    )


# =====================
# CREATE TABLES
# =====================
Base.metadata.create_all(bind=engine)


# =====================
# TOKENIZER
# =====================
tokenizer = tiktoken.get_encoding("cl100k_base")

# =====================
# HELPERS
# =====================
def extract_pdf(path):
    result = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                result.append((i + 1, text))
    return result

def extract_docx(path):
    doc = Document(path)
    return [
        (f"paragraph_{i+1}", p.text)
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]

def chunk_tokens(text, size=500, overlap=80):
    tokens = tokenizer.encode(text)
    chunks = []

    step = size - overlap
    if step <= 0:
        raise ValueError("chunk size must be larger than overlap")

    start = 0
    while start < len(tokens):
        end = start + size
        piece = tokens[start:end]
        chunks.append((tokenizer.decode(piece), len(piece)))
        start += step

    return chunks

# =====================
# API: UPLOAD DOCUMENT
# =====================
@app.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".docx"]:
        raise HTTPException(400, "Only PDF/DOCX allowed")

    stored_name = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, stored_name)

    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    db = SessionLocal()

    try:
        doc = DocumentDB(
            original_filename=file.filename,
            stored_filename=stored_name
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        chunk_idx = 0

        if ext == ".pdf":
            pages = extract_pdf(file_path)
            if not pages:
                raise ValueError("PDF has no extractable text")

            for page, text in pages:
                for content, tc in chunk_tokens(text):
                    db.add(ChunkDB(
                        doc_id=doc.id,
                        chunk_index=chunk_idx,
                        content=content,
                        page=page,
                        token_count=tc
                    ))
                    chunk_idx += 1

        else:  # DOCX
            for section, text in extract_docx(file_path):
                for content, tc in chunk_tokens(text):
                    db.add(ChunkDB(
                        doc_id=doc.id,
                        chunk_index=chunk_idx,
                        content=content,
                        section=section,
                        token_count=tc
                    ))
                    chunk_idx += 1

        db.commit()

        return {
            "doc_id": doc.id,
            "original_filename": doc.original_filename,
            "stored_filename": doc.stored_filename,
            "num_chunks": chunk_idx
        }

    except Exception:
        db.rollback()
        print("🔥 UPLOAD ERROR 🔥")
        traceback.print_exc()
        raise HTTPException(500, "Internal Server Error")

    finally:
        db.close()

# =====================
# API: GET CHUNKS
# =====================
@app.get("/documents/{doc_id}/chunks")
def get_document_chunks(doc_id: int):
    db = SessionLocal()

    try:
        chunks = (
            db.query(ChunkDB)
            .filter(ChunkDB.doc_id == doc_id)
            .order_by(ChunkDB.chunk_index)
            .all()
        )

        if not chunks:
            raise HTTPException(404, "No chunks found")

        return {
            "doc_id": doc_id,
            "num_chunks": len(chunks),
            "chunks": [
                {
                    "chunk_index": c.chunk_index,
                    "content": c.content,
                    "token_count": c.token_count,
                    "page": c.page,
                    "section": c.section
                }
                for c in chunks
            ]
        }

    finally:
        db.close()
# build_vector_db()
@app.get("/rag/search")
def rag_search(q: str, k: int = 3):
    if not vector_ready or len(metadata) == 0:
        raise HTTPException(500, "Vector DB not ready")

    q_vec = embed(q)

    D, I = index.search(
        np.array([q_vec]).astype("float32"), k
    )

    results = []
    for idx in I[0]:
        if idx < len(metadata):
            results.append(metadata[idx])

    return {
        "query": q,
        "results": results
    }

